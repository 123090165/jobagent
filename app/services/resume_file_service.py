from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from app.services.errors import JobAgentError


SUPPORTED_RESUME_FILE_TYPES = {
    ".txt": "txt",
    ".md": "md",
    ".pdf": "pdf",
    ".docx": "docx",
}
DEFAULT_MAX_RESUME_FILE_BYTES = 5 * 1024 * 1024
MAX_RESUME_FILE_BYTES_ENV = "JOBAGENT_MAX_RESUME_FILE_BYTES"
MAX_PDF_PAGES = 50
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024
MAX_DOCX_ARCHIVE_ENTRIES = 2_000


class ResumeFileParseError(JobAgentError):
    """Raised when an uploaded resume file cannot be converted to text."""


def normalize_resume_filename(filename: str | None) -> str:
    normalized = (filename or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
    if not normalized:
        raise ResumeFileParseError("filename is required", "resume_filename_required")
    return normalized


def get_resume_file_type(filename: str | None) -> str:
    normalized = normalize_resume_filename(filename)
    extension = Path(normalized).suffix.lower()
    if extension not in SUPPORTED_RESUME_FILE_TYPES:
        supported = ", ".join(sorted(SUPPORTED_RESUME_FILE_TYPES))
        file_type = extension or "missing extension"
        raise ResumeFileParseError(
            f"unsupported resume file type: {file_type}. supported types: {supported}",
            "resume_file_type_unsupported",
        )
    return SUPPORTED_RESUME_FILE_TYPES[extension]


def get_max_resume_file_bytes() -> int:
    raw_limit = os.getenv(MAX_RESUME_FILE_BYTES_ENV)
    if raw_limit is None or not raw_limit.strip():
        return DEFAULT_MAX_RESUME_FILE_BYTES

    try:
        max_bytes = int(raw_limit)
    except ValueError as exc:
        raise ResumeFileParseError(
            f"{MAX_RESUME_FILE_BYTES_ENV} must be a positive integer",
            "resume_file_size_limit_invalid",
        ) from exc

    if max_bytes <= 0:
        raise ResumeFileParseError(
            f"{MAX_RESUME_FILE_BYTES_ENV} must be a positive integer",
            "resume_file_size_limit_invalid",
        )
    return max_bytes


def validate_resume_file_size(content: bytes) -> None:
    if len(content) > get_max_resume_file_bytes():
        raise ResumeFileParseError("resume file is too large", "resume_file_too_large")


def extract_text_from_resume_file(filename: str | None, content: bytes) -> str:
    """Validate a supported resume file and return extracted plain text."""
    file_type = get_resume_file_type(filename)
    validate_resume_file_size(content)
    if not content:
        raise ResumeFileParseError("resume file cannot be empty", "resume_file_empty")

    if file_type == "pdf":
        return _extract_pdf_text(content)
    if file_type == "docx":
        return _extract_docx_text(content)

    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ResumeFileParseError(
            "resume file must be UTF-8 text",
            "resume_file_decode_failed",
        ) from exc

    extracted_text = text.strip()
    if not extracted_text:
        raise ResumeFileParseError("resume file cannot be empty", "resume_file_empty")
    return extracted_text


def _extract_pdf_text(content: bytes) -> str:
    try:
        import fitz

        with fitz.open(stream=content, filetype="pdf") as document:
            if document.page_count > MAX_PDF_PAGES:
                raise ResumeFileParseError(
                    f"resume PDF exceeds the {MAX_PDF_PAGES}-page limit",
                    "resume_pdf_too_many_pages",
                )
            text = "\n".join(page.get_text("text") for page in document)
    except ResumeFileParseError:
        raise
    except Exception as exc:
        raise ResumeFileParseError(
            "resume PDF could not be read",
            "resume_pdf_parse_failed",
        ) from exc
    return _require_extracted_text(text, file_type="PDF")


def _extract_docx_text(content: bytes) -> str:
    _validate_docx_archive(content)
    try:
        from docx import Document

        document = Document(BytesIO(content))
        lines = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    except Exception as exc:
        raise ResumeFileParseError(
            "resume DOCX could not be read",
            "resume_docx_parse_failed",
        ) from exc
    return _require_extracted_text("\n".join(lines), file_type="DOCX")


def _validate_docx_archive(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise ResumeFileParseError(
                    "resume DOCX contains too many archive entries",
                    "resume_docx_too_complex",
                )
            if sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ResumeFileParseError(
                    "resume DOCX expands beyond the allowed size",
                    "resume_docx_too_complex",
                )
    except ResumeFileParseError:
        raise
    except BadZipFile as exc:
        raise ResumeFileParseError(
            "resume DOCX could not be read",
            "resume_docx_parse_failed",
        ) from exc


def _require_extracted_text(text: str, *, file_type: str) -> str:
    extracted_text = text.strip()
    if not extracted_text:
        raise ResumeFileParseError(
            f"resume {file_type} contains no extractable text; scanned files require OCR",
            "resume_file_no_extractable_text",
        )
    return extracted_text
