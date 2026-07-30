"""回归验证resume intake api的正常链路、失败边界和兼容契约。"""

from __future__ import annotations

from io import BytesIO

import fitz
from docx import Document
from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)


def test_post_resume_text_stores_resume_and_updates_session(
    monkeypatch,
    tmp_path,
) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-text.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()

    response = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-text",
        json={"text": "Python backend engineer\nFastAPI\nSQL"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["resume_document"]["session_id"] == session["session_id"]
    assert payload["resume_document"]["source_type"] == "text"
    assert payload["resume_document"]["text_length"] == len("Python backend engineer\nFastAPI\nSQL")
    assert payload["profile_session"]["resume_document_id"] == payload["resume_document"]["resume_document_id"]
    assert payload["profile_session"]["current_step"] == "resume_ready"
    assert payload["profile_session"]["parsed_review_id"] is None
    assert payload["profile_session"]["profile_draft_id"] is None
    assert payload["profile_session"]["confirmed_profile_id"] is None


def test_post_resume_text_empty_returns_resume_empty(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-empty.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()

    response = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-text",
        json={"text": "   "},
    )

    assert response.status_code == 400
    assert response.json()["error_code"] == "resume_empty"


def test_post_resume_file_accepts_txt(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-file-txt.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()

    response = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-file",
        files={"file": ("resume.txt", b"Resume text", "text/plain")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["resume_document"]["filename"] == "resume.txt"
    assert payload["resume_document"]["file_type"] == "txt"
    assert payload["profile_session"]["current_step"] == "resume_ready"


def test_post_resume_file_accepts_md(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-file-md.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()

    response = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-file",
        files={"file": ("resume.md", b"# Resume", "text/markdown")},
    )

    assert response.status_code == 200
    assert response.json()["resume_document"]["file_type"] == "md"


def test_post_resume_file_accepts_pdf(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-file-pdf.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Backend Engineer\nPython FastAPI SQL")
    content = document.tobytes()
    document.close()

    response = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-file",
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert response.status_code == 200
    assert response.json()["resume_document"]["file_type"] == "pdf"
    assert "Backend Engineer" in response.json()["resume_document"]["text"]


def test_post_resume_file_accepts_docx_and_continues_to_parse(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-file-docx.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()
    document = Document()
    document.add_paragraph("Backend Engineer")
    document.add_paragraph("Skills: Python, FastAPI, SQL")
    buffer = BytesIO()
    document.save(buffer)

    upload = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-file",
        files={
            "file": (
                "resume.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert upload.status_code == 200
    assert upload.json()["resume_document"]["file_type"] == "docx"
    parse = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/parse-resume",
        params={"use_llm": False},
    )
    assert parse.status_code == 200
    assert "Python" in parse.json()["parsed_review"]["skills"]["items"]


def test_post_resume_file_rejects_legacy_doc(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-file-doc.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()

    response = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-file",
        files={"file": ("resume.doc", b"legacy word", "application/msword")},
    )

    assert response.status_code == 400
    assert response.json()["error_code"] == "resume_file_unsupported_type"


def test_get_resume_returns_current_document(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-get.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()
    created = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-text",
        json={"text": "Stored resume"},
    ).json()

    response = client.get(f"/api/v1/profile-sessions/{session['session_id']}/resume")

    assert response.status_code == 200
    assert response.json()["resume_document_id"] == created["resume_document"]["resume_document_id"]


def test_submitting_new_resume_clears_downstream_ids(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-reset.sqlite3"))
    session = client.post("/api/v1/profile-sessions").json()
    first = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-text",
        json={"text": "First resume"},
    ).json()

    response = client.post(
        f"/api/v1/profile-sessions/{session['session_id']}/resume-text",
        json={"text": "Second resume"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["resume_document"]["resume_document_id"] != first["resume_document"]["resume_document_id"]
    assert payload["profile_session"]["parsed_review_id"] is None
    assert payload["profile_session"]["profile_draft_id"] is None
    assert payload["profile_session"]["confirmed_profile_id"] is None
    assert payload["profile_session"]["current_step"] == "resume_ready"


def test_resume_text_unknown_session_returns_not_found(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("JOBAGENT_DB_PATH", str(tmp_path / "resume-missing-session.sqlite3"))

    response = client.post(
        "/api/v1/profile-sessions/missing-session/resume-text",
        json={"text": "Resume"},
    )

    assert response.status_code == 404
    assert response.json()["error_code"] == "profile_session_not_found"
