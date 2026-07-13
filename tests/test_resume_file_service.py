from __future__ import annotations

from io import BytesIO

import fitz
import pytest
from docx import Document

from app.services.resume_file_service import (
    DEFAULT_MAX_RESUME_FILE_BYTES,
    ResumeFileParseError,
    extract_text_from_resume_file,
    get_max_resume_file_bytes,
    get_resume_file_type,
    normalize_resume_filename,
)

SAMPLE_RESUME = """
Backend Engineer
Skills: Python, FastAPI, SQL
Projects:
JobAgent - Built profile session APIs and resume intake flow.
""".strip()


def test_extract_text_from_txt_file() -> None:
    text = extract_text_from_resume_file("resume.txt", SAMPLE_RESUME.encode("utf-8"))

    assert text == SAMPLE_RESUME.strip()
    assert get_resume_file_type("resume.txt") == "txt"


def test_extract_text_from_md_file() -> None:
    markdown_resume = f"# Resume\n\n{SAMPLE_RESUME}"

    text = extract_text_from_resume_file("resume.md", markdown_resume.encode("utf-8"))

    assert text.startswith("# Resume")
    assert get_resume_file_type("resume.md") == "md"


def test_extract_text_from_pdf_file() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Backend Engineer\nPython FastAPI SQL")
    content = document.tobytes()
    document.close()

    text = extract_text_from_resume_file("resume.pdf", content)

    assert "Backend Engineer" in text
    assert "Python FastAPI SQL" in text
    assert get_resume_file_type("resume.pdf") == "pdf"


def test_extract_text_from_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Backend Engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, FastAPI, SQL"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text_from_resume_file("resume.docx", buffer.getvalue())

    assert "Backend Engineer" in text
    assert "Skills\tPython, FastAPI, SQL" in text
    assert get_resume_file_type("resume.docx") == "docx"


def test_extract_text_normalizes_filename() -> None:
    filename = normalize_resume_filename(r"C:\fake\path\resume.TXT")
    text = extract_text_from_resume_file(filename, SAMPLE_RESUME.encode("utf-8"))

    assert filename == "resume.TXT"
    assert text == SAMPLE_RESUME.strip()


def test_extract_text_rejects_empty_file() -> None:
    with pytest.raises(ResumeFileParseError, match="resume file cannot be empty"):
        extract_text_from_resume_file("resume.txt", b"")

    with pytest.raises(ResumeFileParseError, match="resume file cannot be empty") as exc_info:
        extract_text_from_resume_file("resume.md", b"   \n\t")
    assert exc_info.value.error_code == "resume_file_empty"


def test_extract_text_rejects_unsupported_extension() -> None:
    with pytest.raises(ResumeFileParseError, match="unsupported resume file type") as exc_info:
        extract_text_from_resume_file("resume.doc", b"fake word document")
    assert exc_info.value.error_code == "resume_file_type_unsupported"


def test_extract_text_rejects_pdf_without_text_layer() -> None:
    document = fitz.open()
    document.new_page()
    content = document.tobytes()
    document.close()

    with pytest.raises(ResumeFileParseError, match="no extractable text") as exc_info:
        extract_text_from_resume_file("resume.pdf", content)

    assert exc_info.value.error_code == "resume_file_no_extractable_text"


def test_extract_text_rejects_invalid_docx() -> None:
    with pytest.raises(ResumeFileParseError, match="DOCX could not be read") as exc_info:
        extract_text_from_resume_file("resume.docx", b"not a zip archive")

    assert exc_info.value.error_code == "resume_docx_parse_failed"


def test_extract_text_rejects_decode_failure() -> None:
    with pytest.raises(ResumeFileParseError, match="resume file must be UTF-8 text") as exc_info:
        extract_text_from_resume_file("resume.txt", b"\xff\xfe\x00\x00")
    assert exc_info.value.error_code == "resume_file_decode_failed"


def test_extract_text_accepts_file_under_default_size_limit(monkeypatch) -> None:
    monkeypatch.delenv("JOBAGENT_MAX_RESUME_FILE_BYTES", raising=False)
    content = b"a" * DEFAULT_MAX_RESUME_FILE_BYTES

    assert get_max_resume_file_bytes() == DEFAULT_MAX_RESUME_FILE_BYTES
    text = extract_text_from_resume_file("resume.txt", content)

    assert len(text) == DEFAULT_MAX_RESUME_FILE_BYTES


def test_extract_text_rejects_file_over_default_size_limit(monkeypatch) -> None:
    monkeypatch.delenv("JOBAGENT_MAX_RESUME_FILE_BYTES", raising=False)
    content = b"a" * (DEFAULT_MAX_RESUME_FILE_BYTES + 1)

    with pytest.raises(ResumeFileParseError, match="resume file is too large") as exc_info:
        extract_text_from_resume_file("resume.txt", content)

    assert exc_info.value.error_code == "resume_file_too_large"


def test_extract_text_uses_environment_size_limit(monkeypatch) -> None:
    monkeypatch.setenv("JOBAGENT_MAX_RESUME_FILE_BYTES", "16")

    assert get_max_resume_file_bytes() == 16
    assert extract_text_from_resume_file("resume.md", b"short resume") == "short resume"

    with pytest.raises(ResumeFileParseError, match="resume file is too large") as exc_info:
        extract_text_from_resume_file("resume.md", b"this resume is too long")

    assert exc_info.value.error_code == "resume_file_too_large"
